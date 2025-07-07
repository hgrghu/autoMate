"""
Export Manager Module
导出管理器模块，支持多种格式的文档导出和导入
"""

import os
import json
import html2text
import markdown
from pathlib import Path
from PyQt6.QtWidgets import QMessageBox
from PyQt6.QtGui import QTextDocument, QPrinter, QPageLayout, QPageSize
from PyQt6.QtCore import QMarginsF, Qt
from PyQt6.QtPrintSupport import QPrintDialog

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import weasyprint
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False

class ExportManager:
    """导出管理器"""
    
    def __init__(self, parent=None):
        self.parent = parent
        
    def export_document(self, document, file_path, format_type):
        """导出文档"""
        try:
            if format_type == "html":
                return self.export_html(document, file_path)
            elif format_type == "pdf":
                return self.export_pdf(document, file_path)
            elif format_type == "docx":
                return self.export_docx(document, file_path)
            elif format_type == "md":
                return self.export_markdown(document, file_path)
            elif format_type == "txt":
                return self.export_text(document, file_path)
            else:
                raise ValueError(f"不支持的导出格式: {format_type}")
        except Exception as e:
            if self.parent:
                QMessageBox.critical(self.parent, "导出错误", f"导出失败: {str(e)}")
            return False
            
    def export_html(self, document, file_path):
        """导出为HTML"""
        html_content = document.toHtml()
        
        # 添加CSS样式
        styled_html = f"""
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>文档</title>
    <style>
        body {{
            font-family: "Microsoft YaHei", "Helvetica Neue", Arial, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            background-color: #fff;
        }}
        h1, h2, h3, h4, h5, h6 {{
            color: #2c3e50;
            margin-top: 1.5em;
            margin-bottom: 0.5em;
        }}
        p {{
            margin-bottom: 1em;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 1em 0;
        }}
        table, th, td {{
            border: 1px solid #ddd;
        }}
        th, td {{
            padding: 8px;
            text-align: left;
        }}
        th {{
            background-color: #f2f2f2;
        }}
        img {{
            max-width: 100%;
            height: auto;
            display: block;
            margin: 1em auto;
        }}
        .highlight {{
            background-color: #fff3cd;
            padding: 2px 4px;
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>
"""
        
        with open(file_path, 'w', encoding='utf-8') as file:
            file.write(styled_html)
        return True
        
    def export_pdf(self, document, file_path):
        """导出为PDF"""
        try:
            # 方法1: 使用Qt的QPrinter
            printer = QPrinter(QPrinter.PrinterMode.HighResolution)
            printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
            printer.setOutputFileName(file_path)
            
            # 设置页面格式
            page_layout = QPageLayout()
            page_layout.setPageSize(QPageSize(QPageSize.PageSizeId.A4))
            page_layout.setOrientation(QPageLayout.Orientation.Portrait)
            page_layout.setMargins(QMarginsF(20, 20, 20, 20), QPageLayout.Unit.Millimeter)
            printer.setPageLayout(page_layout)
            
            # 打印文档
            document.print(printer)
            return True
            
        except Exception as e:
            # 方法2: 如果可用，使用weasyprint
            if WEASYPRINT_AVAILABLE:
                try:
                    html_content = document.toHtml()
                    weasyprint.HTML(string=html_content).write_pdf(file_path)
                    return True
                except Exception as e2:
                    pass
            raise e
            
    def export_docx(self, document, file_path):
        """导出为Word文档"""
        if not DOCX_AVAILABLE:
            raise ImportError("需要安装python-docx库才能导出Word文档")
            
        doc = Document()
        
        # 设置文档样式
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(12)
        
        # 转换HTML到纯文本并处理格式
        plain_text = document.toPlainText()
        
        # 简单的段落处理
        paragraphs = plain_text.split('\n')
        for para_text in paragraphs:
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
                
        doc.save(file_path)
        return True
        
    def export_markdown(self, document, file_path):
        """导出为Markdown"""
        html_content = document.toHtml()
        
        # 使用html2text转换
        h = html2text.HTML2Text()
        h.ignore_links = False
        h.ignore_images = False
        h.ignore_tables = False
        h.body_width = 0  # 不换行
        
        markdown_content = h.handle(html_content)
        
        with open(file_path, 'w', encoding='utf-8') as file:
            file.write(markdown_content)
        return True
        
    def export_text(self, document, file_path):
        """导出为纯文本"""
        plain_text = document.toPlainText()
        
        with open(file_path, 'w', encoding='utf-8') as file:
            file.write(plain_text)
        return True
        
    # 导入功能
    def import_docx(self, file_path):
        """导入Word文档"""
        if not DOCX_AVAILABLE:
            raise ImportError("需要安装python-docx库才能导入Word文档")
            
        doc = Document(file_path)
        text_content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
                
        # 简单的HTML格式化
        html_content = "<html><body>"
        for text in text_content:
            html_content += f"<p>{text}</p>"
        html_content += "</body></html>"
        
        return html_content
        
    def import_markdown(self, file_path):
        """导入Markdown文档"""
        with open(file_path, 'r', encoding='utf-8') as file:
            markdown_content = file.read()
            
        # 转换为HTML
        html_content = markdown.markdown(
            markdown_content,
            extensions=['tables', 'fenced_code', 'toc']
        )
        
        return f"<html><body>{html_content}</body></html>"
        
    def import_text(self, file_path):
        """导入纯文本文档"""
        with open(file_path, 'r', encoding='utf-8') as file:
            text_content = file.read()
            
        # 转换为HTML
        lines = text_content.split('\n')
        html_content = "<html><body>"
        for line in lines:
            if line.strip():
                html_content += f"<p>{html.escape(line)}</p>"
            else:
                html_content += "<br>"
        html_content += "</body></html>"
        
        return html_content
        
    def get_supported_export_formats(self):
        """获取支持的导出格式"""
        formats = {
            "html": "HTML 文档 (*.html)",
            "pdf": "PDF 文档 (*.pdf)",
            "md": "Markdown 文档 (*.md)",
            "txt": "纯文本 (*.txt)"
        }
        
        if DOCX_AVAILABLE:
            formats["docx"] = "Word 文档 (*.docx)"
            
        return formats
        
    def get_supported_import_formats(self):
        """获取支持的导入格式"""
        formats = {
            "html": "HTML 文档 (*.html)",
            "md": "Markdown 文档 (*.md)",
            "txt": "纯文本 (*.txt)"
        }
        
        if DOCX_AVAILABLE:
            formats["docx"] = "Word 文档 (*.docx)"
            
        return formats